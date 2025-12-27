"""
اسکریپت تولید گزارش Word برای استاد
این فایل یک گزارش کامل با تمام مراحل و تصاویر ایجاد می‌کند.
"""

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os
from datetime import datetime
import arabic_reshaper

def reshape_arabic_text(text):
    """تبدیل متن فارسی/عربی برای نمایش صحیح در Word"""
    try:
        # فقط reshape می‌کنیم، get_display برای Word لازم نیست
        # چون Word خودش RTL را مدیریت می‌کند
        # reshape باعث اتصال صحیح حروف فارسی می‌شود
        reshaped_text = arabic_reshaper.reshape(text)
        return reshaped_text
    except Exception as e:
        # در صورت خطا، متن اصلی را برمی‌گرداند
        print(f"Warning: Could not reshape text: {e}")
        return text

def add_heading_rtl(doc, text, level=1):
    """افزودن عنوان با راست‌چین"""
    # تبدیل متن برای نمایش صحیح فارسی
    text = reshape_arabic_text(text)
    heading = doc.add_heading(text, level=level)
    heading.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    for run in heading.runs:
        run.font.name = 'B Nazanin'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), 'B Nazanin')
    return heading

def add_paragraph_rtl(doc, text, bold=False, font_size=11):
    """افزودن پاراگراف با راست‌چین"""
    # تبدیل متن برای نمایش صحیح فارسی
    text = reshape_arabic_text(text)
    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = para.add_run(text)
    run.font.name = 'B Nazanin'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), 'B Nazanin')
    run.font.size = Pt(font_size)
    if bold:
        run.bold = True
    return para

def add_image_to_doc(doc, image_path, width=6, caption=""):
    """افزودن تصویر به سند"""
    if os.path.exists(image_path):
        para = doc.add_paragraph()
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = para.add_run()
        run.add_picture(image_path, width=Inches(width))
        
        if caption:
            caption_para = doc.add_paragraph()
            caption_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            # تبدیل متن برای نمایش صحیح فارسی
            caption_reshaped = reshape_arabic_text(caption)
            caption_run = caption_para.add_run(caption_reshaped)
            caption_run.font.name = 'B Nazanin'
            caption_run._element.rPr.rFonts.set(qn('w:eastAsia'), 'B Nazanin')
            caption_run.font.size = Pt(10)
            caption_run.italic = True
    else:
        add_paragraph_rtl(doc, f"⚠ تصویر یافت نشد: {image_path}", font_size=10)

def create_report():
    """ایجاد گزارش کامل"""
    
    # ایجاد سند
    doc = Document()
    
    # تنظیم فونت پیش‌فرض
    style = doc.styles['Normal']
    font = style.font
    font.name = 'B Nazanin'
    font.size = Pt(12)
    
    # صفحه عنوان
    title_para = doc.add_paragraph()
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_text = reshape_arabic_text('گزارش پروژه پردازش سیگنال صوتی')
    title_run = title_para.add_run(title_text)
    title_run.font.name = 'B Nazanin'
    title_run._element.rPr.rFonts.set(qn('w:eastAsia'), 'B Nazanin')
    title_run.font.size = Pt(18)
    title_run.bold = True
    
    doc.add_paragraph()
    
    subtitle_para = doc.add_paragraph()
    subtitle_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle_text = reshape_arabic_text('تشخیص بخش‌های واکدار، بی‌واک و سکوت')
    subtitle_run = subtitle_para.add_run(subtitle_text)
    subtitle_run.font.name = 'B Nazanin'
    subtitle_run._element.rPr.rFonts.set(qn('w:eastAsia'), 'B Nazanin')
    subtitle_run.font.size = Pt(14)
    
    doc.add_paragraph()
    
    date_para = doc.add_paragraph()
    date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    date_text = reshape_arabic_text(f'تاریخ: {datetime.now().strftime("%Y/%m/%d")}')
    date_run = date_para.add_run(date_text)
    date_run.font.name = 'B Nazanin'
    date_run._element.rPr.rFonts.set(qn('w:eastAsia'), 'B Nazanin')
    date_run.font.size = Pt(11)
    
    doc.add_page_break()
    
    # فهرست مطالب
    add_heading_rtl(doc, 'فهرست مطالب', level=1)
    add_paragraph_rtl(doc, '1. مقدمه', bold=True)
    add_paragraph_rtl(doc, '2. بخش 1: دریافت و پردازش اولیه سیگنال صوتی', bold=True)
    add_paragraph_rtl(doc, '   2.1. خواندن و نمایش فایل صوتی', font_size=10)
    add_paragraph_rtl(doc, '   2.2. محاسبه انرژی کوتاه‌مدت', font_size=10)
    add_paragraph_rtl(doc, '3. بخش 2: تشخیص بخش‌های واکدار و بی‌واک و سکوت', bold=True)
    add_paragraph_rtl(doc, '   3.1. تقسیم سیگنال به فریم‌های زمانی کوتاه', font_size=10)
    add_paragraph_rtl(doc, '   3.2. محاسبه ZCR (Zero Crossing Rate)', font_size=10)
    add_paragraph_rtl(doc, '   3.3. طبقه‌بندی بر اساس ZCR و انرژی', font_size=10)
    add_paragraph_rtl(doc, '   3.4. استفاده از اتوکرولیشن', font_size=10)
    add_paragraph_rtl(doc, '   3.5. ترکیب روش‌ها برای بهبود تشخیص', font_size=10)
    add_paragraph_rtl(doc, '4. نتیجه‌گیری', bold=True)
    
    doc.add_page_break()
    
    # مقدمه
    add_heading_rtl(doc, '1. مقدمه', level=1)
    add_paragraph_rtl(doc, 
        'این پروژه به بررسی و تحلیل سیگنال‌های صوتی برای شناسایی بخش‌های مختلف گفتار می‌پردازد. '
        'هدف اصلی، تشخیص سه نوع بخش در یک سیگنال صوتی است:'
    )
    add_paragraph_rtl(doc, '• بخش‌های واکدار (Voiced): بخش‌هایی که دارای تناوب مشخص و فرکانس پایه هستند', font_size=10)
    add_paragraph_rtl(doc, '• بخش‌های بی‌واک (Unvoiced): بخش‌هایی که شبیه نویز هستند و تناوب مشخصی ندارند', font_size=10)
    add_paragraph_rtl(doc, '• بخش‌های سکوت (Silence): بخش‌هایی که فاقد انرژی قابل توجه هستند', font_size=10)
    
    add_paragraph_rtl(doc, 
        'این پروژه با استفاده از تکنیک‌های پردازش سیگنال دیجیتال، ویژگی‌های مختلف سیگنال را استخراج کرده '
        'و با ترکیب این ویژگی‌ها، دقت تشخیص را بهبود می‌بخشد.'
    )
    
    doc.add_page_break()
    
    # بخش 1
    add_heading_rtl(doc, '2. بخش 1: دریافت و پردازش اولیه سیگنال صوتی', level=1)
    
    # 1a
    add_heading_rtl(doc, '2.1. خواندن و نمایش فایل صوتی', level=2)
    add_paragraph_rtl(doc, 
        'در این بخش، فایل صوتی خوانده می‌شود و اطلاعات اولیه آن نمایش داده می‌شود. '
        'سیگنال در دو نمودار نمایش داده می‌شود: نمودار دامنه زمانی و نمودار طیف فرکانسی.'
    )
    
    if os.path.exists('part1a_audio_display.png'):
        add_image_to_doc(doc, 'part1a_audio_display.png', width=6, 
                        caption='شکل 1: نمایش سیگنال صوتی در دامنه زمانی و فرکانسی')
    else:
        add_paragraph_rtl(doc, '⚠ تصویر part1a_audio_display.png یافت نشد. لطفاً ابتدا فایل part1a_read_audio.py را اجرا کنید.', font_size=10)
    
    add_paragraph_rtl(doc, 
        'نمودار بالا نشان می‌دهد که سیگنال چگونه در طول زمان تغییر می‌کند و همچنین '
        'توزیع انرژی در فرکانس‌های مختلف چگونه است.'
    )
    
    # 1b
    add_heading_rtl(doc, '2.2. محاسبه انرژی کوتاه‌مدت', level=2)
    add_paragraph_rtl(doc, 
        'انرژی کوتاه‌مدت یکی از مهم‌ترین ویژگی‌های سیگنال صوتی است که برای تشخیص سکوت و گفتار استفاده می‌شود. '
        'فرمول انرژی کوتاه‌مدت به صورت زیر است:'
    )
    
    formula_para = doc.add_paragraph()
    formula_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    formula_run = formula_para.add_run('E(n) = Σ [x(m)]²')
    formula_run.font.name = 'Times New Roman'
    formula_run.font.size = Pt(12)
    formula_run.italic = True
    
    add_paragraph_rtl(doc, 
        'که در آن x(m) نمونه‌های فریم n هستند. در بخش‌های سکوت، انرژی بسیار پایین است '
        'و در بخش‌های گفتار، انرژی بالاتر است.'
    )
    
    if os.path.exists('part1b_short_term_energy.png'):
        add_image_to_doc(doc, 'part1b_short_term_energy.png', width=6,
                        caption='شکل 2: انرژی کوتاه‌مدت و دامنه RMS')
    else:
        add_paragraph_rtl(doc, '⚠ تصویر part1b_short_term_energy.png یافت نشد.', font_size=10)
    
    doc.add_page_break()
    
    # بخش 2
    add_heading_rtl(doc, '3. بخش 2: تشخیص بخش‌های واکدار و بی‌واک و سکوت', level=1)
    
    # 2a
    add_heading_rtl(doc, '3.1. تقسیم سیگنال به فریم‌های زمانی کوتاه', level=2)
    add_paragraph_rtl(doc, 
        'سیگنال صوتی به فریم‌های کوچک‌تر (20 میلی‌ثانیه) تقسیم می‌شود. این کار به دلایل زیر انجام می‌شود:'
    )
    add_paragraph_rtl(doc, '• فرض ایستایی: در بازه‌های زمانی کوتاه، ویژگی‌های سیگنال تقریباً ثابت هستند', font_size=10)
    add_paragraph_rtl(doc, '• پردازش کارآمد: تحلیل فریم به فریم سریع‌تر و دقیق‌تر است', font_size=10)
    add_paragraph_rtl(doc, '• شناسایی دقیق: می‌توان تغییرات جزئی را در طول زمان تشخیص داد', font_size=10)
    
    if os.path.exists('part2a_frame_segmentation.png'):
        add_image_to_doc(doc, 'part2a_frame_segmentation.png', width=6,
                        caption='شکل 3: تقسیم سیگنال به فریم‌های 20 میلی‌ثانیه‌ای')
    else:
        add_paragraph_rtl(doc, '⚠ تصویر part2a_frame_segmentation.png یافت نشد.', font_size=10)
    
    # 2b
    add_heading_rtl(doc, '3.2. محاسبه ZCR (Zero Crossing Rate)', level=2)
    add_paragraph_rtl(doc, 
        'ZCR یا نرخ عبور از صفر، تعداد تغییرات علامت سیگنال در یک فریم است. '
        'این ویژگی برای تشخیص تفاوت بین واکدار و بی‌واک بسیار مفید است.'
    )
    add_paragraph_rtl(doc, 
        '• بخش‌های واکدار: ZCR پایین (زیرا سیگنال تناوبی است)', font_size=10)
    add_paragraph_rtl(doc, 
        '• بخش‌های بی‌واک: ZCR بالا (زیرا شبیه نویز است)', font_size=10)
    add_paragraph_rtl(doc, 
        '• بخش‌های سکوت: ZCR متوسط (به دلیل نویز پس‌زمینه)', font_size=10)
    
    if os.path.exists('part2b_zcr_calculation.png'):
        add_image_to_doc(doc, 'part2b_zcr_calculation.png', width=6,
                        caption='شکل 4: تغییرات ZCR در طول زمان')
    else:
        add_paragraph_rtl(doc, '⚠ تصویر part2b_zcr_calculation.png یافت نشد.', font_size=10)
    
    # 2c
    add_heading_rtl(doc, '3.3. طبقه‌بندی بر اساس ZCR و انرژی', level=2)
    add_paragraph_rtl(doc, 
        'در این بخش، از ترکیب دو ویژگی ZCR و انرژی برای طبقه‌بندی استفاده می‌شود. '
        'قوانین طبقه‌بندی به صورت زیر است:'
    )
    add_paragraph_rtl(doc, '• سکوت: انرژی < آستانه سکوت', font_size=10)
    add_paragraph_rtl(doc, '• واکدار: انرژی > آستانه واکدار و ZCR < آستانه واکدار', font_size=10)
    add_paragraph_rtl(doc, '• بی‌واک: در غیر این صورت', font_size=10)
    
    if os.path.exists('part2c_classification.png'):
        add_image_to_doc(doc, 'part2c_classification.png', width=6,
                        caption='شکل 5: طبقه‌بندی فریم‌ها بر اساس ZCR و انرژی')
    else:
        add_paragraph_rtl(doc, '⚠ تصویر part2c_classification.png یافت نشد.', font_size=10)
    
    # 2d
    add_heading_rtl(doc, '3.4. استفاده از اتوکرولیشن', level=2)
    add_paragraph_rtl(doc, 
        'اتوکرولیشن یک تابع ریاضی است که شباهت یک سیگنال با نسخه جابجا شده خودش را اندازه‌گیری می‌کند. '
        'در بخش‌های واکدار، اتوکرولیشن قله‌های مشخصی دارد که نشان‌دهنده تناوب است.'
    )
    add_paragraph_rtl(doc, 
        'از اتوکرولیشن می‌توان برای محاسبه فرکانس پایه (F0) استفاده کرد. '
        'فاصله بین قله‌های اصلی اتوکرولیشن معکوس فرکانس پایه است.'
    )
    
    if os.path.exists('part2d_autocorrelation.png'):
        add_image_to_doc(doc, 'part2d_autocorrelation.png', width=6,
                        caption='شکل 6: استفاده از اتوکرولیشن برای تشخیص واکدار')
    else:
        add_paragraph_rtl(doc, '⚠ تصویر part2d_autocorrelation.png یافت نشد.', font_size=10)
    
    # 2e
    add_heading_rtl(doc, '3.5. ترکیب روش‌ها برای بهبود تشخیص', level=2)
    add_paragraph_rtl(doc, 
        'در این بخش، از ترکیب سه ویژگی استفاده می‌شود: انرژی کوتاه‌مدت، ZCR، و قدرت اتوکرولیشن. '
        'این ترکیب منجر به دقت بالاتر و کاهش خطاهای تشخیص می‌شود.'
    )
    
    if os.path.exists('part2e_combined_method.png'):
        add_image_to_doc(doc, 'part2e_combined_method.png', width=6,
                        caption='شکل 7: نتایج روش ترکیبی')
    else:
        add_paragraph_rtl(doc, '⚠ تصویر part2e_combined_method.png یافت نشد.', font_size=10)
    
    if os.path.exists('part2e_final_result.png'):
        add_image_to_doc(doc, 'part2e_final_result.png', width=6,
                        caption='شکل 8: نتیجه نهایی - سیگنال با بخش‌های واکدار (سبز)، بی‌واک (نارنجی) و سکوت (خاکستری)')
    else:
        add_paragraph_rtl(doc, '⚠ تصویر part2e_final_result.png یافت نشد.', font_size=10)
    
    doc.add_page_break()
    
    # نتیجه‌گیری
    add_heading_rtl(doc, '4. نتیجه‌گیری', level=1)
    add_paragraph_rtl(doc, 
        'این پروژه نشان می‌دهد که چگونه می‌توان با استفاده از ویژگی‌های ساده سیگنال صوتی '
        '(انرژی، ZCR، اتوکرولیشن)، بخش‌های مختلف گفتار را شناسایی کرد.'
    )
    add_paragraph_rtl(doc, 
        'ترکیب این روش‌ها منجر به دقت بالاتر و قابلیت اعتماد بیشتر می‌شود. '
        'این پروژه می‌تواند در کاربردهای مختلفی مانند بهبود کیفیت انتقال صوت، '
        'تشخیص گفتار، و تحلیل گفتار استفاده شود.'
    )
    
    # ذخیره فایل
    output_file = 'گزارش_پروژه_پردازش_سیگنال_صوتی.docx'
    doc.save(output_file)
    print(f"✅ گزارش با موفقیت در فایل '{output_file}' ذخیره شد.")
    print("⚠ توجه: برای نمایش صحیح متن فارسی، فونت 'B Nazanin' باید در سیستم شما نصب باشد.")
    print("✅ از کتابخانه arabic-reshaper برای اتصال صحیح حروف فارسی استفاده شده است.")
    
    return output_file

if __name__ == '__main__':
    print("در حال ایجاد گزارش Word...")
    create_report()

